import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil
import subprocess

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

src_docx = r'C:\Users\ashra\Downloads\BurnTheMap_Project_Synopsis_updated__2_.docx'
out_dir = r'C:\Users\ashra\Downloads'
mod_docx = os.path.join(out_dir, 'BurnTheMap_Project_Synopsis_Updated_Final.docx')
mod_pdf = os.path.join(out_dir, 'BurnTheMap_Project_Synopsis_Updated_Final.pdf')

doc = docx.Document(src_docx)

# 1. Update Chapter 1 Reference Numbers (Unbold & reset font style to match paragraph)
for p in doc.paragraphs:
    if any(k in p.text for k in ['Chapter 1', '1.1', '1.2', '1.3', 'Background and Motivation', 'The Genesis', 'Table 1.1']):
        for r in p.runs:
            if r.text.strip() in ('[14]', '[1]-[4], [6]'):
                r.bold = False
                r.font.name = None # inherits paragraph font
                r.font.size = Pt(12)
                print(f"Unbolded reference run: {repr(r.text)}")

# 2. Reformat Section 4.1 Project Objective into points starting with "To develop"
target_p = None
for p in doc.paragraphs:
    if 'The core objective is to design and engineer a full-stack, map-first real estate marketplace. This involves developing' in p.text:
        target_p = p
        break

if target_p is not None:
    print("Found Section 4.1 target paragraph!")
    
    # Intro text
    target_p.text = "The core objective is to design and engineer a full-stack, map-first real estate marketplace through the following key objectives:"
    target_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    points = [
        "To develop a responsive, server-rendered Next.js frontend with an interactive Google Maps search surface, custom pins and marker clustering as the primary discovery interface.",
        "To develop advanced multi-criteria filtering across price, bedrooms, locality and amenities that stays synchronized live with the map view.",
        "To develop a RESTful Node.js/Express backend with full CRUD for listings and real-time synchronization between owner updates and buyer views.",
        "To develop a persistent Saved Properties feature alongside a criteria-driven Search Alert subscription mechanism.",
        "To develop an automated Nodemailer email engine that dispatches alerts the moment a new listing matches saved criteria.",
        "To develop secure JWT authentication and authorization, with OAuth planned as an extension, to distinguish buyer/renter from owner/agent accounts.",
        "To develop unit, integration and system-level testing pipelines covering search accuracy, map responsiveness and alert-delivery latency."
    ]
    
    # Insert bullet paragraphs after target_p
    parent = target_p._element.getparent()
    target_idx = parent.index(target_p._element)
    
    for i, pt_text in enumerate(points):
        p_elem = docx.oxml.OxmlElement('w:p')
        pPr = docx.oxml.OxmlElement('w:pPr')
        
        pStyle = docx.oxml.OxmlElement('w:pStyle')
        pStyle.set(docx.oxml.ns.qn('w:val'), 'BodyText')
        pPr.append(pStyle)
        
        spacing = docx.oxml.OxmlElement('w:spacing')
        spacing.set(docx.oxml.ns.qn('w:line'), '271')
        spacing.set(docx.oxml.ns.qn('w:lineRule'), 'auto')
        spacing.set(docx.oxml.ns.qn('w:after'), '120')
        pPr.append(spacing)
        
        ind = docx.oxml.OxmlElement('w:ind')
        ind.set(docx.oxml.ns.qn('w:left'), '720')
        ind.set(docx.oxml.ns.qn('w:hanging'), '360')
        pPr.append(ind)
        
        jc = docx.oxml.OxmlElement('w:jc')
        jc.set(docx.oxml.ns.qn('w:val'), 'both')
        pPr.append(jc)
        
        p_elem.append(pPr)
        
        # Bullet symbol run
        r_b = docx.oxml.OxmlElement('w:r')
        t_b = docx.oxml.OxmlElement('w:t')
        t_b.set(docx.oxml.ns.qn('xml:space'), 'preserve')
        t_b.text = "• "
        r_b.append(t_b)
        p_elem.append(r_b)
        
        # Text run
        r_t = docx.oxml.OxmlElement('w:r')
        t_t = docx.oxml.OxmlElement('w:t')
        t_t.text = pt_text
        r_t.append(t_t)
        p_elem.append(r_t)
        
        parent.insert(target_idx + 1 + i, p_elem)

# 3. Update Section Margins (Top 1", Bottom 1", Right 1", Left 1.5", Header 0.5", Footer 0.5")
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    
    # Configure page borders XML on section
    sectPr = section._sectPr
    pgBorders = sectPr.find(docx.oxml.ns.qn('w:pgBorders'))
    if pgBorders is None:
        pgBorders = docx.oxml.OxmlElement('w:pgBorders')
        sectPr.append(pgBorders)
    
    pgBorders.set(docx.oxml.ns.qn('w:offsetFrom'), 'text')
    
    for side in ('top', 'left', 'bottom', 'right'):
        side_elem = pgBorders.find(docx.oxml.ns.qn(f'w:{side}'))
        if side_elem is None:
            side_elem = docx.oxml.OxmlElement(f'w:{side}')
            pgBorders.append(side_elem)
        side_elem.set(docx.oxml.ns.qn('w:val'), 'single')
        side_elem.set(docx.oxml.ns.qn('w:sz'), '8')
        side_elem.set(docx.oxml.ns.qn('w:space'), '24')
        side_elem.set(docx.oxml.ns.qn('w:color'), '000000')

doc.save(mod_docx)
print("Saved updated docx via python-docx!")

# 4. Now use MS Word COM in PowerShell to open mod_docx, save, and export PDF cleanly
ps_code = f"""
$docPath = "{mod_docx}"
$pdfPath = "{mod_pdf}"

$word = New-Object -ComObject Word.Application
$word.Visible = $false

$doc = $word.Documents.Open($docPath)

foreach ($sect in $doc.Sections) {{
    $ps = $sect.PageSetup
    $ps.TopMargin = 72      # 1.0 inch
    $ps.BottomMargin = 72   # 1.0 inch
    $ps.LeftMargin = 108    # 1.5 inch
    $ps.RightMargin = 72    # 1.0 inch
    $ps.HeaderDistance = 36 # 0.5 inch
    $ps.FooterDistance = 36 # 0.5 inch
    
    $ps.Borders.DistanceFrom = 1
    $ps.Borders.Enable = 1
}}

$doc.Save()
$doc.ExportAsFixedFormat($pdfPath, 17)
$doc.Close()
$word.Quit()
"""

with open('run_word_export.ps1', 'w', encoding='utf-8') as f:
    f.write(ps_code)

res = subprocess.run(['powershell', '-ExecutionPolicy', 'Bypass', '-File', 'run_word_export.ps1'], capture_output=True, text=True)
print("STDOUT:", res.stdout)
print("STDERR:", res.stderr)

if os.path.exists('run_word_export.ps1'):
    os.remove('run_word_export.ps1')

# Overwrite original updated__2_ (1) (1) (1) files with these freshly created ones
orig_docx = r'C:\Users\ashra\Downloads\BurnTheMap_Project_Synopsis_updated__2_ (1) (1) (1).docx'
orig_pdf = r'C:\Users\ashra\Downloads\BurnTheMap_Project_Synopsis_updated__2_ (1) (1) (1).pdf'

shutil.copyfile(mod_docx, orig_docx)
if os.path.exists(mod_pdf):
    shutil.copyfile(mod_pdf, orig_pdf)
    print("Updated original files in Downloads!")
